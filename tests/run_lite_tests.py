from __future__ import annotations

import json
import os
import shutil
import subprocess
import sys
import zipfile
from pathlib import Path

from docx import Document


REPO_ROOT = Path(__file__).resolve().parent.parent
SANDBOX = REPO_ROOT / "tests" / "lite_sandbox"
SKILL_ROOT = REPO_ROOT / "packages" / "claude" / ".claude" / "skills" / "mathmodel-lite"
PREFLIGHT = SKILL_ROOT / "scripts" / "lite_preflight.py"
RUNNER = SKILL_ROOT / "scripts" / "lite_run.py"
FINALIZER = SKILL_ROOT / "scripts" / "lite_finalize.py"


def run(script: Path, cwd: Path, *args: str) -> subprocess.CompletedProcess[str]:
    return subprocess.run(
        [sys.executable, str(script), *args],
        cwd=str(cwd),
        text=True,
        encoding="utf-8",
        errors="replace",
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        check=False,
    )


def write_json(path: Path, data: object) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def assert_true(condition: bool, message: str) -> None:
    if not condition:
        raise AssertionError(message)


def prepare(name: str) -> Path:
    cwd = SANDBOX / name
    if cwd.exists():
        shutil.rmtree(cwd)
    problem = cwd / "problem_files"
    problem.mkdir(parents=True)
    (problem / "problem.txt").write_text("Q1: fit a linear model and report the slope.\n", encoding="utf-8")
    (problem / "data.csv").write_text("x,y\n1,2\n2,4\n3,6\n", encoding="utf-8")
    completed = run(PREFLIGHT, cwd)
    assert_true(completed.returncode == 0, f"preflight failed\n{completed.stdout}")
    output = cwd / "paper_output_lite"
    write_json(
        output / "plan.json",
        {"delivery": {"mode": "smoke-test", "reason": "Synthetic regression fixture"}, "questions": [{"id": "Q1", "task": "estimate slope", "model": "least squares", "output": "slope"}]},
    )
    model_source = '''from pathlib import Path
import json

root = Path.cwd()
output = root / "paper_output_lite"
table = output / "tables" / "q1.csv"
table.write_text("metric,value\\nslope,2.0\\n", encoding="utf-8")
result = {
    "status": "computed",
    "questions": [{
        "id": "Q1",
        "answer": "The fitted slope is 2.0.",
        "method": "least squares",
        "metrics": {"slope": 2.0},
        "evidence": ["paper_output_lite/tables/q1.csv"]
    }]
}
(output / "results.json").write_text(json.dumps(result, indent=2), encoding="utf-8")
'''
    (output / "code" / "model.py").write_text(model_source, encoding="utf-8")
    (output / "paper.md").write_text(
        "\n".join(
            [
                "# Linear Modeling Paper",
                "# 摘要",
                "本文针对 Q1 建立线性模型并计算斜率。",
                "# 1 问题重述",
                "Q1 要求根据数据估计线性关系。",
                "# 2 模型假设",
                "观测误差独立且数据口径一致。",
                "# 3 模型建立与求解",
                "## Q1 方法、结果与检验",
                "Q1 使用最小二乘模型，斜率为 2.0。通过逐点代入原始数据验证拟合关系；这个构造算例没有测量噪声，不能据此宣称真实样本也有相同表现。",
                "# 4 结果与检验",
                "Q1 的斜率为 2.0，结果与原始数据一致。",
                "# 5 模型评价",
                "模型简单且可解释。",
                "# 6 结论",
                "Q1 已得到明确数值答案。",
            ]
        )
        + "\n",
        encoding="utf-8",
    )
    return cwd


def test_lite_happy_path() -> None:
    cwd = prepare("happy")
    completed = run(RUNNER, cwd)
    assert_true(completed.returncode == 0, f"runner failed\n{completed.stdout}")
    completed = run(FINALIZER, cwd)
    assert_true(completed.returncode == 0, f"finalizer failed\n{completed.stdout}")
    report = json.loads((cwd / "paper_output_lite" / "lite_report.json").read_text(encoding="utf-8"))
    assert_true(report["status"] == "PASS", "Lite report should pass")
    docx_path = cwd / "paper_output_lite" / "paper.docx"
    assert_true(docx_path.stat().st_size > 0, "Lite DOCX should be nonempty")
    docx_text = "\n".join(paragraph.text for paragraph in Document(docx_path).paragraphs)
    assert_true("Q1" in docx_text and "2.0" in docx_text, "Lite DOCX should preserve the computed result")


def test_lite_rejects_modified_input() -> None:
    cwd = prepare("modified_input")
    (cwd / "problem_files" / "data.csv").write_text("x,y\n1,3\n", encoding="utf-8")
    completed = run(RUNNER, cwd)
    assert_true(completed.returncode != 0, "runner should reject modified input")


def test_lite_rejects_modified_model_after_run() -> None:
    cwd = prepare("modified_model")
    assert_true(run(RUNNER, cwd).returncode == 0, "runner should pass before modification")
    model = cwd / "paper_output_lite" / "code" / "model.py"
    model.write_text(model.read_text(encoding="utf-8") + "\n# changed\n", encoding="utf-8")
    completed = run(FINALIZER, cwd)
    assert_true(completed.returncode != 0, "finalizer should reject modified model")


def test_lite_rejects_placeholder_paper() -> None:
    cwd = prepare("placeholder")
    assert_true(run(RUNNER, cwd).returncode == 0, "runner should pass")
    paper = cwd / "paper_output_lite" / "paper.md"
    paper.write_text(paper.read_text(encoding="utf-8") + "\n待补结果。\n", encoding="utf-8")
    completed = run(FINALIZER, cwd)
    assert_true(completed.returncode != 0, "finalizer should reject placeholder paper")


def test_lite_preflight_rejects_other_editions_in_all_skill_roots() -> None:
    cases = [
        ("skills", "paper-workflow-orchestrator", None),
        (".agents/skills", "pro-workflow-orchestrator", None),
        (".codex/skills", "custom-standard", "standard"),
        (".claude/skills", "custom-pro", "pro"),
        (".trae/skills", "paper-workflow-orchestrator", None),
    ]
    for index, (root_text, entry_name, marker_edition) in enumerate(cases, start=1):
        cwd = SANDBOX / f"lite_rejects_foreign_{index}"
        problem = cwd / "problem_files"
        problem.mkdir(parents=True)
        (problem / "problem.txt").write_text("Q1: test isolation.\n", encoding="utf-8")
        entry = cwd / Path(root_text) / entry_name
        entry.mkdir(parents=True)
        if marker_edition:
            write_json(entry / "MATHMODEL_EDITION.json", {
                "product": "MathModel-Skill",
                "edition": marker_edition,
                "version": "test",
                "entry_skill": entry_name,
            })
        else:
            (entry / "SKILL.md").write_text(f"---\nname: {entry_name}\n---\n", encoding="utf-8")
        completed = run(PREFLIGHT, cwd)
        assert_true(completed.returncode != 0, f"Lite should reject {root_text}/{entry_name}")
        manifest = json.loads((cwd / "paper_output_lite" / "input_manifest.json").read_text(encoding="utf-8"))
        assert_true(manifest["mathmodel_installation"]["detected_editions"], "foreign edition should be recorded")
        assert_true(any("当前入口是 Lite" in item or "混装" in item for item in manifest["failures"]), "mixed-edition failure should be explicit")


def test_lite_release_packages_are_deterministic_and_marked() -> None:
    builder = REPO_ROOT / "scripts" / "build_release_packages.py"
    first = SANDBOX / "release-first"
    second = SANDBOX / "release-second"
    for output in (first, second):
        completed = subprocess.run(
            [sys.executable, str(builder), "--output-dir", str(output)],
            cwd=str(REPO_ROOT),
            text=True,
            encoding="utf-8",
            errors="replace",
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
            check=False,
        )
        assert_true(completed.returncode == 0, completed.stdout)
    names = sorted(path.name for path in first.glob("*.zip"))
    assert_true(names == [
        "MathModel-Skill-Lite-Claude-Code.zip",
        "MathModel-Skill-Lite-Codex.zip",
        "MathModel-Skill-Lite-Trae.zip",
    ], "Lite should build exactly three platform archives")
    assert_true((first / "SHA256SUMS.txt").read_bytes() == (second / "SHA256SUMS.txt").read_bytes(), "checksum files should match")
    for name in names:
        assert_true((first / name).read_bytes() == (second / name).read_bytes(), f"nondeterministic archive: {name}")
        with zipfile.ZipFile(first / name) as archive:
            entries = archive.namelist()
            marker = next(item for item in entries if item.endswith("mathmodel-lite/MATHMODEL_EDITION.json"))
            marker_data = json.loads(archive.read(marker).decode("utf-8"))
            assert_true(marker_data["edition"] == "lite", "Lite marker edition should be present")
            assert_true(marker_data["version"] == (REPO_ROOT / "VERSION").read_text().strip(), "Lite marker version should match")
            assert_true("AGENTS.md" not in entries and "CLAUDE.md" not in entries, "archives must not overwrite user instructions")
            if "Codex" in name:
                assert_true(marker.startswith(".agents/skills/"), "Codex must install in the modern skill root")
            assert_true("LICENSE" in entries, "Lite archive should include MIT license")


def test_lite_body_and_numeric_coverage() -> None:
    for name, edit in (
        ("missing_q_heading", lambda t: t.replace("## Q1 方法、结果与检验", "## 方法")),
        ("wrong_number", lambda t: t.replace("2.0", "9999")),
        ("missing_body", lambda t: "# 摘要\n# Q1\n2.0\n# 问题重述\n# 假设\n# 模型\n# 检验\n# 评价\n# 结论\n"),
    ):
        cwd = prepare(name)
        assert_true(run(RUNNER, cwd).returncode == 0, "runner should pass")
        paper = cwd / "paper_output_lite/paper.md"
        paper.write_text(edit(paper.read_text(encoding="utf-8")), encoding="utf-8")
        assert_true(run(FINALIZER, cwd).returncode != 0, f"must reject {name}")


def test_lite_default_scope_is_not_smoke() -> None:
    cwd = prepare("default_scope")
    plan = cwd / "paper_output_lite/plan.json"
    data = json.loads(plan.read_text(encoding="utf-8"))
    data.pop("delivery")
    write_json(plan, data)
    assert_true(run(RUNNER, cwd).returncode == 0, "run should pass")
    assert_true(run(FINALIZER, cwd).returncode != 0, "a smoke fixture is not a basic report")


def test_lite_plan_and_input_inventory_invalidation() -> None:
    for name in ("plan", "added_input"):
        cwd = prepare("changed_" + name)
        assert_true(run(RUNNER, cwd).returncode == 0, "run should pass")
        if name == "plan":
            path = cwd / "paper_output_lite/plan.json"
            data = json.loads(path.read_text(encoding="utf-8"))
            data["questions"][0]["task"] = "Changed question"
            write_json(path, data)
        else:
            (cwd / "problem_files/new.txt").write_text("new input", encoding="utf-8")
        assert_true(run(FINALIZER, cwd).returncode != 0, "stale inputs must invalidate finalization")


def test_lite_noop_does_not_reuse_results() -> None:
    cwd = prepare("noop")
    assert_true(run(RUNNER, cwd).returncode == 0, "first run should pass")
    (cwd / "paper_output_lite/code/model.py").write_text("pass\n", encoding="utf-8")
    assert_true(run(RUNNER, cwd).returncode != 0, "no-op run must fail")
    assert_true(run(FINALIZER, cwd).returncode != 0, "old results cannot be delivered")


def test_lite_timeout_records_failure() -> None:
    cwd = prepare("timeout")
    (cwd / "paper_output_lite/code/model.py").write_text("import time\ntime.sleep(30)\n", encoding="utf-8")
    assert_true(run(RUNNER, cwd, "--timeout", "1").returncode != 0, "watchdog must stop the model")
    report = json.loads((cwd / "paper_output_lite/run_manifest.json").read_text(encoding="utf-8"))
    assert_true(report["status"] == "FAIL" and "exceeded" in report["failures"][0], "timeout must leave a failed receipt")


def test_lite_rejects_escaping_paths() -> None:
    sys.path.insert(0, str(SKILL_ROOT / "scripts"))
    from lite_common import safe_path
    for value in ("../outside", "/tmp/outside", "C:outside", "C:/outside", "a\\..\\outside"):
        try:
            safe_path(REPO_ROOT, value)
        except ValueError:
            continue
        raise AssertionError(f"path accepted: {value}")


def test_lite_duplicate_result_questions_rejected() -> None:
    cwd = prepare("duplicate_questions")
    model = cwd / "paper_output_lite/code/model.py"
    text = model.read_text(encoding="utf-8").replace('(output / "results.json").write_text', 'result["questions"] *= 2\n(output / "results.json").write_text')
    model.write_text(text, encoding="utf-8")
    assert_true(run(RUNNER, cwd).returncode == 0, "run records the actual duplicate results")
    assert_true(run(FINALIZER, cwd).returncode != 0, "duplicate IDs must fail")


def test_lite_docx_contains_actual_image() -> None:
    cwd = prepare("image")
    model = cwd / "paper_output_lite/code/model.py"
    model.write_text(model.read_text(encoding="utf-8") + '\nimport matplotlib\nmatplotlib.use("Agg")\nimport matplotlib.pyplot as plt\nplt.plot([1,2,3], [2,4,6])\nplt.savefig(output / "figures/q1.png")\n', encoding="utf-8")
    assert_true(run(RUNNER, cwd).returncode == 0, "image run should pass")
    paper = cwd / "paper_output_lite/paper.md"
    paper.write_text(paper.read_text(encoding="utf-8") + "\n![Linear fit](paper_output_lite/figures/q1.png)\n", encoding="utf-8")
    completed = run(FINALIZER, cwd)
    assert_true(completed.returncode == 0, completed.stdout)
    assert_true(len(Document(cwd / "paper_output_lite/paper.docx").inline_shapes) == 1, "image must be embedded, not replaced with a caption")


def test_lite_non_utf8_host_console() -> None:
    cwd = prepare("western_console")
    model = cwd / "paper_output_lite/code/model.py"
    model.write_text(model.read_text(encoding="utf-8") + '\nprint("中文计算日志")\n', encoding="utf-8")
    env = {**os.environ, "PYTHONUTF8": "0", "PYTHONIOENCODING": "cp1252"}
    for script in (RUNNER, FINALIZER):
        result = subprocess.run([sys.executable, str(script)], cwd=cwd, env=env, capture_output=True, encoding="utf-8", errors="replace")
        assert_true(result.returncode == 0, result.stdout + result.stderr)


def main() -> int:
    if SANDBOX.exists():
        shutil.rmtree(SANDBOX)
    SANDBOX.mkdir(parents=True)
    tests = [
        test_lite_happy_path,
        test_lite_rejects_modified_input,
        test_lite_rejects_modified_model_after_run,
        test_lite_rejects_placeholder_paper,
        test_lite_body_and_numeric_coverage,
        test_lite_default_scope_is_not_smoke,
        test_lite_plan_and_input_inventory_invalidation,
        test_lite_noop_does_not_reuse_results,
        test_lite_timeout_records_failure,
        test_lite_rejects_escaping_paths,
        test_lite_duplicate_result_questions_rejected,
        test_lite_docx_contains_actual_image,
        test_lite_non_utf8_host_console,
        test_lite_preflight_rejects_other_editions_in_all_skill_roots,
        test_lite_release_packages_are_deterministic_and_marked,
    ]
    for test in tests:
        test()
        print(f"[PASS] {test.__name__}")
    print("All Lite tests passed.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
