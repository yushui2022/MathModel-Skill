from __future__ import annotations

import copy
import importlib.util
import json
import os
import shutil
import sys
import tempfile
import unittest
import zipfile
from pathlib import Path
from unittest import mock

from pro_fixture import (
    REPO, SKILLS, SCRIPTS, prepare, complete, complete_consensus, approve,
    candidates, tournament, run, envelope, write_test_reviews,
)
from pro_contracts import (
    safe_path, output_root, contract, read_json, sha256_file, write_json, validate_envelope,
)
from pro_checkpoint import require_checkpoints, validate_checkpoint_artifacts
from pro_validation import (
    check_sources, check_tournament, receipts, check_replication, check_robustness,
    check_review, check_freeze, compare, check_claims,
)
from pro_format_check import text_coverage, check_paper
import pro_preflight
import pro_render_pdf
from test_authoring_policy import AuthoringPolicyTests

TEMP = Path(os.environ.get("MATHMODEL_TEST_TEMP", tempfile.gettempdir())).resolve()
TEMP.mkdir(parents=True, exist_ok=True)


class CoreTests(unittest.TestCase):
    def setUp(self):
        self.project = Path(tempfile.mkdtemp(prefix="pro-test-", dir=TEMP))
        self.root = prepare(self.project)

    def tearDown(self):
        if self.project.is_relative_to(TEMP) and self.project.name.startswith("pro-test-"):
            shutil.rmtree(self.project)

    def test_skill_surface_is_pro_only(self):
        names = {p.name for p in SKILLS.iterdir() if (p / "SKILL.md").is_file()}
        self.assertEqual(len(names), 10)
        self.assertIn("pro-workflow-orchestrator", names)
        self.assertNotIn("paper-micro-unit-generator", names)
        self.assertFalse((REPO / "packages/trae").exists())

    def test_forward_evaluation_has_four_required_case_types(self):
        cases = [read_json(p) for p in (REPO / "tests/fixtures/forward").glob("*/case.json")]
        self.assertEqual({c["category"] for c in cases}, {"prediction", "optimization", "graph", "open_data"})
        self.assertTrue((REPO / "tests/fixtures/forward/optimization/solve_enumeration.py").is_file())

    def test_latest_frontier_model_profiles_are_selected(self):
        catalog = pro_preflight.load_model_catalog()
        for model in ("gpt-6-astra", "Claude Fable 5.1", "Claude Opus 5", "Claude Sonnet 5", "Claude Fable 5", "gpt-5.6-sol ultra"):
            with self.subTest(model=model):
                profile, _ = pro_preflight.match_model_profile(model, catalog)
                self.assertIsNotNone(profile)
                self.assertTrue(pro_preflight.reasoning_profile(profile, "ultra")["compatible"])
        self.assertIsNone(pro_preflight.match_model_profile("unknown-future-model", catalog)[0])

    def test_preflight_warns_ordinary_model_and_blocks_mixed_install(self):
        result = run(SCRIPTS / "pro_preflight.py", "--project-root", self.project, "--platform", "codex", "--model", "ordinary")
        self.assertIn("WARNING", result.stdout)
        for skill_root in ("skills", ".agents/skills", ".codex/skills", ".claude/skills", ".trae/skills"):
            with self.subTest(skill_root=skill_root):
                path = self.project / skill_root / "mathmodel-lite"
                path.mkdir(parents=True)
                (path / "SKILL.md").write_text("legacy entry", encoding="utf-8")
                result = run(SCRIPTS / "pro_preflight.py", "--project-root", self.project, "--platform", "codex", "--model", "astra", check=False)
                self.assertNotEqual(result.returncode, 0)
                shutil.rmtree(path)

    def test_old_pro_installation_is_blocked(self):
        path = self.project / ".agents/skills/pro-workflow-orchestrator"
        path.mkdir(parents=True)
        write_json(path / "MATHMODEL_EDITION.json", {"edition": "pro", "version": "3.0.0"})
        result = run(SCRIPTS / "pro_preflight.py", "--project-root", self.project, "--platform", "codex", "--model", "astra", check=False)
        self.assertNotEqual(result.returncode, 0)

    def test_checkpoints_cannot_skip_and_hash_change_invalidates_downstream(self):
        result = run(SCRIPTS / "pro_checkpoint.py", "approve", "--checkpoint", "2", "--decision", "fixture", "--project-root", self.project, check=False)
        self.assertNotEqual(result.returncode, 0)
        approve(self.project, 1)
        data = read_json(self.root / "problem_consensus.json")
        data["consensus"].append("change")
        write_json(self.root / "problem_consensus.json", data)
        self.assertTrue(require_checkpoints(self.project, self.root, 1))
        self.assertEqual(read_json(self.root / "checkpoint_ledger.json")["checkpoints"]["1"]["status"], "PENDING")

    def test_explicit_decision_is_required(self):
        result = run(SCRIPTS / "pro_checkpoint.py", "approve", "--checkpoint", "1", "--project-root", self.project, check=False)
        self.assertNotEqual(result.returncode, 0)

    def test_pending_checkpoint_validation_is_not_success(self):
        result = run(SCRIPTS / "pro_checkpoint.py", "validate", "--checkpoint", "1", "--project-root", self.project, check=False)
        self.assertNotEqual(result.returncode, 0)
        approve(self.project, 1)
        run(SCRIPTS / "pro_checkpoint.py", "validate", "--checkpoint", "1", "--project-root", self.project)

    def test_instruction_audit_is_required_and_new_instruction_invalidates_checkpoint(self):
        approve(self.project, 1)
        (self.project / "AGENTS.md").write_text("Additional instruction", encoding="utf-8")
        self.assertTrue(require_checkpoints(self.project, self.root, 1))

    def test_checkpoint_rejection_invalidates_later_states(self):
        approve(self.project, 1)
        run(SCRIPTS / "pro_checkpoint.py", "reject", "--checkpoint", "1", "--decision", "fixture rejected", "--project-root", self.project)
        self.assertEqual(read_json(self.root / "checkpoint_ledger.json")["checkpoints"]["1"]["status"], "REJECTED")
        self.assertTrue(require_checkpoints(self.project, self.root, 1))

    def test_idempotent_preflight_keeps_approval(self):
        approve(self.project, 1)
        before = sha256_file(self.root / "pro_config.json")
        with mock.patch.object(pro_preflight, "contract", side_effect=lambda **kw: {**contract(**kw), "created_at_utc": "2099-01-01T00:00:00Z"}):
            with mock.patch.object(sys, "argv", ["pro_preflight.py", "--project-root", str(self.project), "--platform", "codex", "--model", "gpt-6-astra", "--reasoning", "ultra"]):
                self.assertEqual(pro_preflight.main(), 0)
        self.assertEqual(before, sha256_file(self.root / "pro_config.json"))
        self.assertFalse(require_checkpoints(self.project, self.root, 1))

    def test_tournament_requires_four_quality_routes_by_default_contract(self):
        self.assertFalse(check_tournament(candidates(), tournament()))
        data = candidates()
        data["subproblems"][0]["routes"] = data["subproblems"][0]["routes"][:2]
        self.assertTrue(check_tournament(data, tournament()))

    def test_tournament_coverage_weights_and_duplicate_ids(self):
        data = candidates()
        data["weights"]["task_fit"] = 9
        self.assertTrue(check_tournament(data, tournament()))
        data = candidates()
        data["subproblems"][0]["routes"][1]["route_id"] = "r1"
        with self.assertRaises(ValueError):
            check_tournament(data, tournament())

    def test_source_gate_rejects_invalid_and_single_publisher_critical_claim(self):
        self.assertTrue(check_sources(self.root, {"sources": [], "critical_claims": [{"claim_id": "C", "source_ids": ["missing"]}]}))
        source = {"source_id": "S1", "url": "https://missing-source.invalid", "title": "x", "publisher": "x",
                  "purpose": "test", "claim_ids": ["C"], "access_status": "PUBLIC_OK", "authorization_required": False,
                  "accessed_at_utc": "invalid", "content_sha256": "not-a-hash"}
        with self.assertRaises(ValueError):
            check_sources(self.root, {"sources": [source]})

    def test_public_capture_rejects_private_addresses(self):
        from pro_capture_source import public_url
        with mock.patch("socket.getaddrinfo", return_value=[(None, None, None, None, ("127.0.0.1", 80))]):
            with self.assertRaises(ValueError):
                public_url("https://example.test/a")

    def test_source_receipt_and_snapshot_are_both_verified(self):
        snapshot = self.root / "research/content.bin"
        snapshot.write_bytes(b"Recorded public source fixture")
        receipt = self.root / "research/retrieval.json"
        timestamp = "2026-01-01T00:00:00Z"
        write_json(receipt, envelope("pro-public-source-capture", url="https://example.org/data", http_status=200,
                                    accessed_at_utc=timestamp, snapshot_path="research/content.bin", content_sha256=sha256_file(snapshot)))
        source = {"source_id": "S1", "url": "https://example.org/data", "title": "fixture", "publisher": "fixture",
                  "purpose": "test", "claim_ids": ["C"], "access_status": "PUBLIC_OK", "authorization_required": False,
                  "accessed_at_utc": timestamp, "snapshot_path": "research/content.bin", "content_sha256": sha256_file(snapshot),
                  "retrieval_receipt": "research/retrieval.json", "retrieval_receipt_sha256": sha256_file(receipt)}
        data = {"sources": [source], "critical_claims": [{"claim_id": "C", "source_ids": ["S1"], "cross_validation_required": False, "single_source_reason": "constructed fixture"}]}
        self.assertFalse(check_sources(self.root, data))
        snapshot.write_bytes(b"changed")
        self.assertTrue(check_sources(self.root, data))
        data["critical_claims"][0]["cross_validation_required"] = True
        self.assertTrue(any("two independent" in e for e in check_sources(self.root, data)))

    def test_public_capture_saves_real_response_without_ambient_credentials(self):
        from pro_capture_source import capture
        response = mock.MagicMock(status_code=200, headers={"Content-Type": "text/plain"})
        response.iter_content.return_value = [b"actual public response fixture"]
        with mock.patch("pro_capture_source.public_url"), mock.patch("pro_capture_source.requests.Session") as factory:
            session = factory.return_value.__enter__.return_value
            session.get.return_value = response
            path = capture(self.root, "https://example.test/data", "Sretrieved")
            self.assertFalse(session.trust_env)
        receipt = read_json(path)
        self.assertEqual((self.root / receipt["snapshot_path"]).read_bytes(), b"actual public response fixture")
        self.assertEqual(receipt["content_sha256"], sha256_file(self.root / receipt["snapshot_path"]))

    def test_public_capture_rejects_unsuccessful_http(self):
        from pro_capture_source import capture
        import requests
        response = mock.MagicMock(status_code=404)
        response.raise_for_status.side_effect = requests.HTTPError("404")
        with mock.patch("pro_capture_source.public_url"), mock.patch("pro_capture_source.requests.Session") as factory:
            factory.return_value.__enter__.return_value.get.return_value = response
            with self.assertRaises(requests.HTTPError):
                capture(self.root, "https://example.test/missing", "Smissing")
        self.assertFalse((self.root / "research/Smissing/retrieval.json").exists())

    def test_path_resolution_handles_posix_windows_mixed_and_chinese(self):
        for value in ("figures/中文.png", "figures\\中文.png", "figures\\子目录/图.png"):
            self.assertTrue(safe_path(self.root, value).is_relative_to(self.root))
        for value in ("/etc/passwd", "C:\\secret.txt", "../outside", "a\\..\\outside", "C:secret", "\\\\host\\share", ""):
            with self.subTest(value=value), self.assertRaises(ValueError):
                safe_path(self.root, value)
        with self.assertRaises(ValueError):
            output_root(self.project, self.project.parent)

    def test_duplicate_json_keys_and_nonfinite_numbers_are_rejected(self):
        path = self.root / "qa/bad.json"
        for text in ('{"x":1,"x":2}', '{"x":NaN}', '{"x":Infinity}', '{"x":1e999}'):
            path.write_text(text, encoding="utf-8")
            with self.assertRaises(ValueError):
                read_json(path)

    def test_missing_libreoffice_is_detected(self):
        with mock.patch.object(pro_render_pdf.shutil, "which", return_value=None), mock.patch.object(Path, "is_file", return_value=False):
            self.assertIsNone(pro_render_pdf.find_soffice(None))

    def test_data_pipeline_resolves_helpers_from_its_own_installation(self):
        path = SKILLS / "data-cleaning-and-visualization/scripts/run_pipeline.py"
        spec = importlib.util.spec_from_file_location("tested_data_pipeline", path)
        module = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(module)
        with mock.patch.object(module.os, "chdir"), mock.patch.object(module.subprocess, "run", return_value=mock.Mock(returncode=0)) as commands:
            self.assertEqual(module.main(), 0)
        self.assertTrue(commands.call_args_list)
        self.assertTrue(all(Path(call.args[0][1]).parent == path.parent for call in commands.call_args_list))

    def test_failure_success_resets_consecutive_counter(self):
        script = SKILLS / "context-memory-keeper/scripts/update_pro_memory.py"
        for failed in (True, False, True, False, True, False):
            extra = ["--failure", "solver failed"] if failed else []
            run(script, "--project-root", self.project, "--phase", "P3", "--next-action", "continue", *extra)
            self.assertFalse(read_json(self.root / "context/workflow_memory.json")["blocked"])
        for _ in range(3):
            run(script, "--project-root", self.project, "--phase", "P3", "--next-action", "continue", "--failure", "solver failed")
        self.assertTrue(read_json(self.root / "context/workflow_memory.json")["blocked"])

    def test_platform_sync_and_deterministic_two_package_build(self):
        run(REPO / "scripts/sync_platform_packages.py", "--check")
        first, second = self.project / "first", self.project / "second"
        for path in (first, second):
            run(REPO / "scripts/build_release_packages.py", "--output-dir", path)
        for path in first.glob("*.zip"):
            self.assertEqual(sha256_file(path), sha256_file(second / path.name))
            with zipfile.ZipFile(path) as archive:
                names = archive.namelist()
                self.assertNotIn("AGENTS.md", names)
                self.assertNotIn("CLAUDE.md", names)
                self.assertFalse(any(n.startswith(".trae/") for n in names))
                self.assertTrue(any("pro_run_experiment.py" in n for n in names))
                self.assertIn("assets/orlando-liu-social.jpg", names)
                self.assertTrue(all(i.create_system == 3 and i.compress_type == zipfile.ZIP_STORED for i in archive.infolist()))
                installed = self.project / "installed" / path.stem
                archive.extractall(installed)
            shutil.copytree(self.project / "problem_files", installed / "problem_files")
            platform = "codex" if "Codex" in path.name else "claude-code"
            skill_root = ".agents" if platform == "codex" else ".claude"
            entry = installed / skill_root / "skills/pro-workflow-orchestrator/scripts/pro_preflight.py"
            self.assertEqual(entry.read_text(encoding="utf-8"), (SCRIPTS / "pro_preflight.py").read_text(encoding="utf-8"))
            run(entry, "--project-root", installed, "--platform", platform, "--model", "gpt-6-astra")
            config = read_json(installed / "paper_output_pro/pro_config.json")
            self.assertEqual(len(config["mathmodel_installation"]["installations"]), 1)
            mixed = installed / ".claude/skills/mathmodel-lite"
            mixed.mkdir(parents=True)
            (mixed / "SKILL.md").write_text("legacy lite entry", encoding="utf-8")
            self.assertNotEqual(run(entry, "--project-root", installed, "--platform", platform, "--model", "gpt-6-astra", check=False).returncode, 0)

    def test_project_location_binding_prevents_copied_approvals(self):
        approve(self.project, 1)
        other = self.project / "copied-project"
        shutil.copytree(self.project / "paper_output_pro", other / "paper_output_pro")
        shutil.copytree(self.project / "problem_files", other / "problem_files")
        self.assertTrue(require_checkpoints(other, other / "paper_output_pro", 1))


@unittest.skipUnless(os.environ.get("REQUIRE_LIBREOFFICE") == "1", "full pipeline requires LibreOffice")
class PipelineTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.project = Path(tempfile.mkdtemp(prefix="pro-pipeline-", dir=TEMP))
        cls.root = complete(cls.project)
        cls.original = {p.relative_to(cls.project): p.read_bytes() for p in cls.project.rglob("*") if p.is_file()}

    def tearDown(self):
        for path in self.project.rglob("*"):
            if path.is_file() and path.relative_to(self.project) not in self.original:
                path.unlink()
        for relative, data in self.original.items():
            path = self.project / relative
            if not path.is_file() or path.read_bytes() != data:
                path.parent.mkdir(parents=True, exist_ok=True)
                path.write_bytes(data)

    @classmethod
    def tearDownClass(cls):
        if cls.project.is_relative_to(TEMP) and cls.project.name.startswith("pro-pipeline-"):
            shutil.rmtree(cls.project)

    def test_complete_synthetic_pro_pipeline_passes_final_gate(self):
        result = run(SCRIPTS / "pro_gate.py", "--project-root", self.project)
        self.assertIn("[PASS]", result.stdout)
        values = [read_json(self.root / "experiments" / n / "metrics.json")["metrics"]["cost"] for n in ("base-milp", "base-enum")]
        self.assertAlmostEqual(*values)
        self.assertEqual(read_json(self.root / "pro_gate_report.json")["acceptance_scope"], "ENGINEERING_SMOKE_ONLY")

    def test_scope_change_invalidates_checkpoints_and_final_gate(self):
        config = read_json(self.root / "pro_config.json")
        from pro_authoring_policy import make_policy
        config["paper_delivery"] = make_policy()
        write_json(self.root / "pro_config.json", config)
        self.assertNotEqual(run(SCRIPTS / "pro_gate.py", "--project-root", self.project, check=False).returncode, 0)
        self.assertEqual(read_json(self.root / "checkpoint_ledger.json")["checkpoints"]["1"]["status"], "PENDING")

    def test_full_paper_review_requires_per_question_assessment(self):
        from pro_authoring_policy import make_policy
        config = read_json(self.root / "pro_config.json")
        config["paper_delivery"] = make_policy()
        write_json(self.root / "pro_config.json", config)
        write_test_reviews(self.root)
        errors = check_review(self.root, read_json(self.root / "review_board_report.json"))
        self.assertTrue(any("subproblem_id" in e for e in errors))

    def test_missing_config_overwrites_old_final_pass_with_blocked(self):
        (self.root / "pro_config.json").unlink()
        self.assertNotEqual(run(SCRIPTS / "pro_gate.py", "--project-root", self.project, check=False).returncode, 0)
        self.assertEqual(read_json(self.root / "pro_gate_report.json")["acceptance_scope"], "NOT_ACCEPTED")

    def test_competition_abstract_and_appendix_have_real_docx_page_boundaries(self):
        from pro_format_check import expected_docx
        from lxml import etree
        import io
        plan = read_json(self.root / "paper_plan.json")
        plan["delivery_mode"] = "competition"
        plan["sections"][-1]["kind"] = "appendix"
        write_json(self.root / "paper_plan.json", plan)
        with zipfile.ZipFile(io.BytesIO(expected_docx(self.root))) as archive:
            xml = etree.fromstring(archive.read("word/document.xml"))
            ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            titles = xml.xpath("//w:p[w:pPr/w:pageBreakBefore]//w:t/text()", namespaces=ns)
            self.assertEqual(titles, ["Model", "Conclusion"])

    def test_libreoffice_renders_docx_when_required_by_ci(self):
        render = read_json(self.root / "render_manifest.json")
        self.assertTrue(render["libreoffice_version"])
        self.assertGreaterEqual(len(render["pages"]), 1)
        self.assertTrue(all((self.root / p["path"]).is_file() for p in render["pages"]))

    def test_raw_input_change_blocks_delivery(self):
        (self.project / "problem_files/data.json").write_text('{"changed":true}', encoding="utf-8")
        self.assertNotEqual(run(SCRIPTS / "pro_gate.py", "--project-root", self.project, check=False).returncode, 0)
        self.assertEqual(read_json(self.root / "checkpoint_ledger.json")["checkpoints"]["1"]["status"], "PENDING")

    def test_added_and_removed_inputs_invalidate_approval(self):
        (self.project / "problem_files/extra.txt").write_text("new task", encoding="utf-8")
        self.assertTrue(require_checkpoints(self.project, self.root, 3))

    def test_new_frozen_file_is_rejected(self):
        (self.root / "code/extra.py").write_text("print(9999)", encoding="utf-8")
        self.assertTrue(check_freeze(self.root, read_json(self.root / "evidence_freeze.json")))

    def test_replication_robustness_and_document_corruption_are_blocking(self):
        runs, errors = receipts(self.root, read_json(self.root / "experiment_manifest.json"))
        self.assertFalse(errors)
        replication = read_json(self.root / "replication_report.json")
        replication["critical_results"][0]["replication_paths"][1]["run_id"] = "baseline-enum"
        self.assertTrue(check_replication(self.root, replication, runs, read_json(self.root / "tournament_report.json")))
        self.assertFalse(compare(1, 9999, {"kind": "exact"}))
        robustness = read_json(self.root / "robustness_report.json")
        robustness["baseline_comparisons"][0]["measurements"][0]["value"] = 9999
        self.assertTrue(check_robustness(self.root, robustness, runs))
        (self.root / "final_paper.docx").write_bytes(b"broken")
        self.assertNotEqual(run(SCRIPTS / "pro_gate.py", "--project-root", self.project, check=False).returncode, 0)

    def test_false_random_statistics_and_repeated_seeds_rejected(self):
        runs, _ = receipts(self.root, read_json(self.root / "experiment_manifest.json"))
        data = read_json(self.root / "robustness_report.json")
        data["stochastic_methods"] = [{"run_ids": ["base-milp"] * 10, "mean": "not measured", "variance": -1}]
        self.assertTrue(check_robustness(self.root, data, runs))

    def test_unknown_claim_evidence_is_rejected(self):
        runs, _ = receipts(self.root, read_json(self.root / "experiment_manifest.json"))
        claims = read_json(self.root / "claim_evidence_map.json")
        claims["claims"][0]["evidence_ids"] = ["invented"]
        self.assertTrue(check_claims(self.root, claims, runs, read_json(self.root / "replication_report.json"), read_json(self.root / "source_ledger.json")))

    def test_changed_result_file_is_rejected(self):
        (self.root / "experiments/base-milp/metrics.json").write_text('{"metrics":{"cost":9999}}', encoding="utf-8")
        self.assertTrue(receipts(self.root, read_json(self.root / "experiment_manifest.json"))[1])

    def test_claim_cannot_borrow_an_unrelated_replication(self):
        runs, _ = receipts(self.root, read_json(self.root / "experiment_manifest.json"))
        claims = read_json(self.root / "claim_evidence_map.json")
        claims["claims"][0]["numeric_evidence"] = [{"run_id": "baseline-milp", "metric": "cost", "decimals": 2, "display": "166.00"}]
        self.assertTrue(check_claims(self.root, claims, runs, read_json(self.root / "replication_report.json"), read_json(self.root / "source_ledger.json")))

    def test_qualitative_source_claim_does_not_need_a_fake_number(self):
        runs, _ = receipts(self.root, read_json(self.root / "experiment_manifest.json"))
        claims = read_json(self.root / "claim_evidence_map.json")
        claims["claims"].append({"claim_id": "Ctheory", "statement": "A cited qualitative assumption", "section_id": "model",
            "claim_type": "qualitative", "qualitative_rationale": "Source-backed scope definition, not a measured estimate",
            "external": True, "source_ids": ["S1"], "evidence_ids": [], "numeric_evidence": []})
        sources = {"sources": [{"source_id": "S1", "claim_ids": ["Ctheory"]}],
                   "critical_claims": [{"claim_id": "Ctheory", "source_ids": ["S1"]}]}
        self.assertFalse(check_claims(self.root, claims, runs, read_json(self.root / "replication_report.json"), sources))

    def test_failed_receipts_cannot_be_omitted(self):
        path = self.root / "experiments/failed"
        path.mkdir(exist_ok=True)
        write_json(path / "receipt.json", envelope("fixture-failed-receipt", run_id="failed"))
        self.assertTrue(receipts(self.root, read_json(self.root / "experiment_manifest.json"))[1])

    def test_review_board_requires_all_roles_and_no_unresolved_major(self):
        report = read_json(self.root / "review_board_report.json")
        self.assertFalse(check_review(self.root, report))
        entry = report["rounds"][-1]["reviews"][0]
        path = self.root / entry["report_path"]
        detail = read_json(path)
        detail["findings"] = [{"finding_id": "M1", "severity": "MAJOR", "evidence": "counterexample", "disposition": "OPEN"}]
        write_json(path, detail)
        entry["report_sha256"] = sha256_file(path)
        self.assertTrue(check_review(self.root, report))

    def test_stale_review_after_manuscript_change_is_rejected(self):
        path = self.root / "final_paper_source.md"
        path.write_text(path.read_text(encoding="utf-8") + "\nChanged conclusion.\n", encoding="utf-8")
        self.assertTrue(check_review(self.root, read_json(self.root / "review_board_report.json")))

    def test_shared_review_context_is_rejected(self):
        report = read_json(self.root / "review_board_report.json")
        for entry in report["rounds"][-1]["reviews"]:
            path = self.root / entry["report_path"]
            detail = read_json(path)
            detail["execution"]["context_id"] = "same"
            write_json(path, detail)
            entry["report_sha256"] = sha256_file(path)
        self.assertTrue(check_review(self.root, report))

    def test_truncated_document_does_not_score_full_coverage(self):
        text = (self.root / "final_paper_source.md").read_text(encoding="utf-8")
        coverage, _ = text_coverage(text, text[:230])
        self.assertLess(coverage, 0.2)

    def test_missing_claim_number_and_placeholder_text_are_rejected(self):
        path = self.root / "final_paper_source.md"
        text = path.read_text(encoding="utf-8")
        value = read_json(self.root / "claim_evidence_map.json")["claims"][0]["numeric_evidence"][0]["display"]
        path.write_text(text.replace(value, "999999") + "\nTODO\n", encoding="utf-8")
        errors = check_paper(self.root)
        self.assertTrue(any("computed number" in e for e in errors))
        self.assertTrue(any("placeholder" in e for e in errors))

    def test_deleted_section_is_rejected(self):
        path = self.root / "final_paper_source.md"
        path.write_text(path.read_text(encoding="utf-8").replace("## Limitations", "## Something else"), encoding="utf-8")
        self.assertTrue(check_paper(self.root))

    def test_missing_visual_page_or_stale_image_is_rejected(self):
        path = self.root / "visual_review.json"
        visual = read_json(path)
        visual["pages"] = []
        write_json(path, visual)
        self.assertNotEqual(run(SCRIPTS / "pro_format_check.py", "--project-root", self.project, check=False).returncode, 0)

    def test_manual_docx_change_is_rejected(self):
        from docx import Document
        path = self.root / "final_paper.docx"
        document = Document(path)
        document.add_paragraph("An unreviewed extra result is 9999.")
        document.save(path)
        self.assertNotEqual(run(SCRIPTS / "pro_format_check.py", "--project-root", self.project, check=False).returncode, 0)

    def test_run_spec_mutation_is_detected(self):
        path = self.root / "code/spec-base-milp.json"
        spec = read_json(path)
        spec["args"].append("--all-open")
        write_json(path, spec)
        self.assertTrue(receipts(self.root, read_json(self.root / "experiment_manifest.json"))[1])

    def test_unapproved_run_route_and_invalid_watchdog_do_not_start(self):
        from pro_run_experiment import execute
        path = self.root / "code/spec-invalid.json"
        spec = read_json(self.root / "code/spec-base-milp.json")
        spec.update(run_id="invalid", route_id="invented")
        write_json(path, spec)
        with self.assertRaises(ValueError):
            execute(self.project, path)
        spec.update(route_id="r2", timeout_seconds=0)
        write_json(path, spec)
        with self.assertRaises(ValueError):
            execute(self.project, path)
        self.assertFalse((self.root / "experiments/invalid").exists())

    def test_unclosed_formula_is_rejected(self):
        path = self.root / "final_paper_source.md"
        path.write_text(path.read_text(encoding="utf-8") + "\n\n$$ x=1\n", encoding="utf-8")
        self.assertTrue(any("unmatched math" in e for e in check_paper(self.root)))

    def test_unreviewed_docx_style_change_is_rejected(self):
        from docx import Document
        from docx.shared import RGBColor
        path = self.root / "final_paper.docx"
        document = Document(path)
        document.styles["Normal"].font.color.rgb = RGBColor(255, 255, 255)
        document.save(path)
        self.assertNotEqual(run(SCRIPTS / "pro_format_check.py", "--project-root", self.project, check=False).returncode, 0)

    def test_external_critical_claim_cannot_skip_cross_validation(self):
        runs, _ = receipts(self.root, read_json(self.root / "experiment_manifest.json"))
        claims = read_json(self.root / "claim_evidence_map.json")
        claims["claims"][0].update(external=True, source_ids=["S1"])
        sources = {"sources": [{"source_id": "S1", "claim_ids": [claims["claims"][0]["claim_id"]]}], "critical_claims": []}
        errors = check_claims(self.root, claims, runs, read_json(self.root / "replication_report.json"), sources)
        self.assertTrue(any("cross-validation" in e for e in errors))

    def test_repeated_substantive_body_is_rejected(self):
        path = self.root / "final_paper_source.md"
        body = path.read_text(encoding="utf-8")
        paragraph = next(p for p in body.split("\n\n") if len(p) > 200)
        path.write_text(body + "\n\n" + paragraph, encoding="utf-8")
        self.assertTrue(any("duplicat" in e for e in check_paper(self.root)))

    def test_multiline_bracket_formula_remains_native(self):
        from pro_format_check import expected_docx
        from lxml import etree
        import io
        path = self.root / "final_paper_source.md"
        body = path.read_text(encoding="utf-8")
        path.write_text(body + "\n\n\\[\nx^2+y^2=1\n\\]\n", encoding="utf-8")
        self.assertFalse(check_paper(self.root))
        with zipfile.ZipFile(io.BytesIO(expected_docx(self.root))) as archive:
            xml = etree.fromstring(archive.read("word/document.xml"))
            self.assertEqual(len(xml.xpath("//*[local-name()='oMath']")), 2)

    def test_compact_table_and_equation_caption_keep_together(self):
        from pro_format_check import expected_docx
        from lxml import etree
        import io
        path = self.root / "final_paper_source.md"
        body = path.read_text(encoding="utf-8")
        path.write_text(body + "\n\nTable 1 Values\n\n| A | B |\n| --- | --- |\n| 1 | 2 |\n\n$$x=1$$\n\nEquation 2 Identity\n", encoding="utf-8")
        with zipfile.ZipFile(io.BytesIO(expected_docx(self.root))) as archive:
            xml = etree.fromstring(archive.read("word/document.xml"))
            ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            self.assertTrue(xml.xpath("//w:tbl/w:tr[1]//w:keepNext", namespaces=ns))
            self.assertTrue(xml.xpath("//w:p[w:r/w:t='Table 1 Values']/w:pPr/w:keepNext", namespaces=ns))

    def test_review_requests_cannot_self_approve(self):
        run(SCRIPTS / "pro_collect_reviews.py", "--project-root", self.project, "--round", "2", "--prepare")
        result = run(SCRIPTS / "pro_collect_reviews.py", "--project-root", self.project, "--round", "2", check=False)
        self.assertNotEqual(result.returncode, 0)
        self.assertEqual(read_json(self.root / "review_board_report.json")["status"], "BLOCKED")


if __name__ == "__main__":
    suite = unittest.TestSuite([
        unittest.defaultTestLoader.loadTestsFromTestCase(CoreTests),
        unittest.defaultTestLoader.loadTestsFromTestCase(AuthoringPolicyTests),
        unittest.defaultTestLoader.loadTestsFromTestCase(PipelineTests),
    ])
    result = unittest.TextTestRunner(verbosity=2).run(suite)
    raise SystemExit(0 if result.wasSuccessful() else 1)
