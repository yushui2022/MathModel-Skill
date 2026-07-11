import argparse
import hashlib
import json
import re
import shutil
import subprocess
import sys
import tempfile
import zipfile
from datetime import datetime
from difflib import SequenceMatcher
from pathlib import Path
from typing import Any

from docx import Document
from lxml import etree

from formula_omml import source_formula_tokens


BASE_DIR = Path.cwd()
OUTPUT_DIR = BASE_DIR / "paper_output"
SOURCE_FILE = OUTPUT_DIR / "final_paper_source.md"
FALLBACK_SOURCE_FILE = OUTPUT_DIR / "final_paper.md"
DOCX_FILE = OUTPUT_DIR / "final_paper.docx"
OUTLINE_FILE = OUTPUT_DIR / "plan" / "paper_outline.json"
FIGURE_INDEX_FILE = OUTPUT_DIR / "figure_index.json"
TABLE_INDEX_FILE = OUTPUT_DIR / "tables" / "table_index.json"
EVIDENCE_GATE_REPORT = OUTPUT_DIR / "qa" / "evidence_gate_report.json"
REPORT_MD = OUTPUT_DIR / "format_check_report.md"
REPORT_JSON = OUTPUT_DIR / "format_check_report.json"
RENDER_DIR = OUTPUT_DIR / "qa" / "rendered"

PLACEHOLDERS = [
    "内容生成中",
    "关键词1",
    "论文题目缺失",
    "TODO",
    "待补",
]

REQUIRED_SECTIONS = [
    "摘要",
    "关键词",
    "1 问题重述",
    "2 问题分析",
    "3 模型假设",
    "4 符号说明",
    "5 模型的建立与求解",
    "6 模型检验",
    "7 模型评价",
    "8 参考文献",
    "附录",
]

CONTROL_CHARACTER_RE = re.compile(r"[\x00-\x09\x0b\x0c\x0e-\x1f\x7f]")

INTERNAL_PROJECT_PATTERNS = [
    ("paper_output 路径", re.compile(r"paper_output(?:[\\/]|\b)", flags=re.IGNORECASE)),
    ("evidence_gate.py", re.compile(r"evidence_gate\.py", flags=re.IGNORECASE)),
    ("workflow_guard.py", re.compile(r"workflow_guard\.py", flags=re.IGNORECASE)),
    ("preflight_check.py", re.compile(r"preflight_check\.py", flags=re.IGNORECASE)),
    ("format_formal_docx.py", re.compile(r"format_formal_docx\.py", flags=re.IGNORECASE)),
    ("check_paper_format.py", re.compile(r"check_paper_format\.py", flags=re.IGNORECASE)),
    ("run_modeling.py", re.compile(r"run_modeling\.py", flags=re.IGNORECASE)),
    ("本样例", re.compile(r"本样例")),
    ("一键草稿/生成", re.compile(r"一键(?:草稿|生成)")),
    ("Skill/Agent 调用话术", re.compile(r"(?:本|该)?\s*(?:skill|agent)\s*(?:调用|生成|负责|读取)", flags=re.IGNORECASE)),
]


def configure_utf8_stdio() -> None:
    for stream in (sys.stdout, sys.stderr):
        try:
            stream.reconfigure(encoding="utf-8")
        except Exception:
            pass


def load_json(path: Path) -> Any:
    if not path.exists():
        return {}
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception as exc:
        return {"__error__": str(exc)}


def rel(path: Path) -> str:
    try:
        return path.relative_to(BASE_DIR).as_posix()
    except ValueError:
        return path.as_posix()


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def source_path() -> Path:
    if SOURCE_FILE.exists():
        return SOURCE_FILE
    return FALLBACK_SOURCE_FILE


def heading_offset(text: str, title_pattern: str) -> int | None:
    match = re.search(rf"(?mi)^\s*#{{1,6}}\s*(?:\d+(?:\.\d+)*\s*)?{title_pattern}\s*$", text)
    return match.start() if match else None


def line_number(text: str, offset: int) -> int:
    return text.count("\n", 0, offset) + 1


def control_character_failures(text: str) -> list[str]:
    failures: list[str] = []
    for match in CONTROL_CHARACTER_RE.finditer(text):
        codepoint = ord(match.group(0))
        failures.append(f"源文件第 {line_number(text, match.start())} 行含控制字符 U+{codepoint:04X}，疑似 LaTeX 反斜杠被错误转义。")
        if len(failures) >= 12:
            break
    return failures


def internal_language_failures(text: str) -> list[str]:
    appendix_start = heading_offset(text, "附录")
    searchable = text[:appendix_start] if appendix_start is not None else text
    failures: list[str] = []
    for label, pattern in INTERNAL_PROJECT_PATTERNS:
        match = pattern.search(searchable)
        if match:
            failures.append(f"正文第 {line_number(text, match.start())} 行存在内部工程话术：{label}。此类实现说明只能放在附录。")
    return failures


def _expand_citation_group(group: str) -> set[int]:
    values: set[int] = set()
    for part in re.split(r"[,，]", group):
        part = part.strip()
        range_match = re.fullmatch(r"(\d+)\s*[-–—]\s*(\d+)", part)
        if range_match:
            start, end = int(range_match.group(1)), int(range_match.group(2))
            if 0 < start <= end <= start + 50:
                values.update(range(start, end + 1))
            continue
        if part.isdigit():
            values.add(int(part))
    return values


def citation_quality(text: str) -> tuple[dict[str, Any], list[str], list[str]]:
    failures: list[str] = []
    warnings: list[str] = []
    references_start = heading_offset(text, "参考文献")
    appendix_start = heading_offset(text, "附录")
    if references_start is None:
        return {
            "bibliography_entry_count": 0,
            "body_citation_count": 0,
            "body_citation_ids": [],
            "uncited_bibliography_ids": [],
            "unknown_body_citation_ids": [],
        }, failures, warnings

    body = text[:references_start]
    references_end = appendix_start if appendix_start is not None and appendix_start > references_start else len(text)
    references = text[references_start:references_end]
    bibliography_ids = {
        int(value)
        for value in re.findall(r"(?m)^\s*[\[［](\d+)[\]］]", references)
    }
    citation_groups = re.findall(r"[\[［](\d+(?:\s*[-–—,，]\s*\d+)*)[\]］]", body)
    body_ids: set[int] = set()
    for group in citation_groups:
        body_ids.update(_expand_citation_group(group))

    if len(bibliography_ids) < 3:
        failures.append(f"参考文献条目不足：{len(bibliography_ids)} < 3。")
    elif len(bibliography_ids) < 5:
        warnings.append(f"参考文献条目少于推荐值：{len(bibliography_ids)} < 5。")
    if bibliography_ids and not citation_groups:
        failures.append("正文没有文献引用；仅在文末列出参考文献不能通过正式门禁。")

    unknown = sorted(body_ids - bibliography_ids)
    uncited = sorted(bibliography_ids - body_ids)
    if unknown:
        failures.append(f"正文引用了参考文献表中不存在的编号：{unknown}")
    if uncited:
        failures.append(f"参考文献表存在未在正文引用的条目：{uncited}")

    return {
        "bibliography_entry_count": len(bibliography_ids),
        "body_citation_count": len(citation_groups),
        "body_citation_ids": sorted(body_ids),
        "uncited_bibliography_ids": uncited,
        "unknown_body_citation_ids": unknown,
    }, failures, warnings


def paragraph_candidates(text: str) -> list[tuple[int, str]]:
    references_start = heading_offset(text, "参考文献")
    body = text[:references_start] if references_start is not None else text
    candidates: list[tuple[int, str]] = []
    in_code = False
    in_formula = False
    for number, line in enumerate(body.splitlines(), start=1):
        stripped = line.strip()
        if stripped.startswith("```"):
            in_code = not in_code
            continue
        if in_code:
            continue
        if stripped.startswith("$$"):
            if stripped.count("$$") < 2:
                in_formula = not in_formula
            continue
        if in_formula or not stripped:
            continue
        if re.match(r"^#{1,6}\s+", stripped):
            continue
        if stripped.startswith(("|", "[[FIGURE:", "[[TABLE:", "![")):
            continue
        if re.match(r"^[-*]\s+", stripped):
            continue
        candidates.append((number, stripped))
    return candidates


def normalize_paragraph(text: str) -> str:
    value = re.sub(r"\\\(.+?\\\)|(?<!\$)\$(?!\$).+?(?<!\\)\$(?!\$)", "<math>", text)
    value = re.sub(r"\d+(?:\.\d+)?%?", "<num>", value)
    value = re.sub(r"[\W_]+", "", value.lower(), flags=re.UNICODE)
    return value


def duplicate_paragraph_failures(text: str) -> tuple[list[str], list[dict[str, Any]]]:
    normalized = [
        (line, paragraph, normalize_paragraph(paragraph))
        for line, paragraph in paragraph_candidates(text)
    ]
    normalized = [item for item in normalized if len(item[2]) >= 45]
    failures: list[str] = []
    matches: list[dict[str, Any]] = []
    for index, (line_a, paragraph_a, norm_a) in enumerate(normalized):
        for line_b, paragraph_b, norm_b in normalized[index + 1:]:
            length_ratio = min(len(norm_a), len(norm_b)) / max(len(norm_a), len(norm_b))
            if length_ratio < 0.82:
                continue
            ratio = 1.0 if norm_a == norm_b else SequenceMatcher(None, norm_a, norm_b, autojunk=False).ratio()
            if ratio < 0.94:
                continue
            match = {
                "line_a": line_a,
                "line_b": line_b,
                "similarity": round(ratio, 4),
                "preview_a": paragraph_a[:100],
                "preview_b": paragraph_b[:100],
            }
            matches.append(match)
            failures.append(f"检测到数字归一化后的重复段落：第 {line_a} 行与第 {line_b} 行，相似度 {ratio:.1%}。")
            if len(failures) >= 20:
                return failures, matches
    return failures, matches


def compact_text(text: str) -> str:
    text = re.sub(r"!\[[^\]]*\]\([^)]+\)", "", text)
    text = re.sub(r"\[[^\]]+\]\([^)]+\)", "", text)
    text = re.sub(r"`{1,3}", "", text)
    text = re.sub(r"[#>*_\-|$`{}\[\]():;,.，。；：！？、\s]", "", text)
    return text


def char_count(text: str) -> dict[str, int]:
    cjk = len(re.findall(r"[\u4e00-\u9fff]", text))
    nonspace = len(re.sub(r"\s+", "", text))
    content = len(compact_text(text))
    return {"cjk": cjk, "nonspace": nonspace, "content": content}


def has_required_section(text: str, label: str) -> bool:
    if label in ("摘要", "关键词", "附录"):
        return re.search(rf"(^|\n)\s*(?:#+\s*)?(?:\*\*)?{re.escape(label)}", text) is not None
    number, title = label.split(" ", 1)
    return re.search(rf"(^|\n)\s*#*\s*{re.escape(number)}\s+.*{re.escape(title)}", text) is not None


def natural_q_key(qid: str) -> tuple[int, str]:
    match = re.search(r"\d+", qid)
    if match:
        return (int(match.group()), qid)
    return (10_000, qid)


def qids_from_outline(outline: Any) -> list[str]:
    if isinstance(outline, dict) and isinstance(outline.get("questions"), list):
        qids = [str(item.get("question_id") or "").strip() for item in outline["questions"] if isinstance(item, dict)]
        qids = [qid for qid in qids if qid]
        if qids:
            return sorted(set(qids), key=natural_q_key)
    model_route = load_json(OUTPUT_DIR / "plan" / "model_route.json")
    qids = []
    for item in model_route.get("questions", []) if isinstance(model_route, dict) else []:
        if isinstance(item, dict) and item.get("question_id"):
            qids.append(str(item["question_id"]))
    return sorted(set(qids), key=natural_q_key)


def index_items(data: Any, key: str, id_key: str) -> list[dict[str, Any]]:
    if not isinstance(data, dict):
        return []
    return [item for item in data.get(key, []) if isinstance(item, dict) and item.get(id_key)]


def referenced(text: str, item: dict[str, Any], id_key: str) -> bool:
    candidates = [
        str(item.get(id_key) or ""),
        str(item.get("title") or ""),
        Path(str(item.get("path") or item.get("expected_path") or "")).stem,
    ]
    candidates = [candidate.strip() for candidate in candidates if candidate and candidate.strip()]
    return any(candidate in text for candidate in candidates)


def check_docx_structure(path: Path) -> dict[str, Any]:
    if not path.exists():
        return {"exists": False}
    try:
        if not zipfile.is_zipfile(path):
            return {"exists": True, "error": "DOCX is not a valid zip package"}
        doc = Document(str(path))
        with zipfile.ZipFile(path) as archive:
            document_xml = archive.read("word/document.xml")
        root = etree.fromstring(document_xml)
        namespaces = {
            "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
        }
        native_omml_count = len(root.xpath(".//m:oMath", namespaces=namespaces))
        display_omml_count = len(root.xpath(".//m:oMathPara", namespaces=namespaces))
        math_text = "".join(root.xpath(".//m:t/text()", namespaces=namespaces))
        headings = [
            paragraph.text
            for paragraph in doc.paragraphs
            if paragraph.style and paragraph.style.name.startswith("Heading") and paragraph.text.strip()
        ]
        text_parts = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
        docx_text = "\n".join(text_parts)
        image_count = 0
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                image_count += 1
        empty_table_count = 0
        table_cell_count = 0
        for table in doc.tables:
            has_text = False
            for row in table.rows:
                for cell in row.cells:
                    table_cell_count += 1
                    if cell.text.strip():
                        has_text = True
            if not has_text:
                empty_table_count += 1
        return {
            "exists": True,
            "package_ok": True,
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
            "table_cell_count": table_cell_count,
            "empty_table_count": empty_table_count,
            "image_count": image_count,
            "inline_shape_count": len(doc.inline_shapes),
            "heading_count": len(headings),
            "sample_headings": headings[:12],
            "nonspace_text_chars": len(re.sub(r"\s+", "", docx_text + math_text)),
            "native_omml_count": native_omml_count,
            "display_omml_count": display_omml_count,
            "math_text_chars": len(re.sub(r"\s+", "", math_text)),
            "text_preview": docx_text[:240],
        }
    except Exception as exc:
        return {"exists": True, "error": str(exc)}


def markdown_heading_count(text: str) -> int:
    return len(re.findall(r"(^|\n)\s*#{1,6}\s+\S+", text))


def visual_qa_failures(
    docx_structure: dict[str, Any],
    source_heading_count: int,
    figure_count: int,
    table_count: int,
) -> tuple[list[str], list[str]]:
    failures: list[str] = []
    warnings: list[str] = []
    if not docx_structure.get("exists") or docx_structure.get("error"):
        return failures, warnings

    paragraph_count = int(docx_structure.get("paragraph_count") or 0)
    heading_count = int(docx_structure.get("heading_count") or 0)
    docx_table_count = int(docx_structure.get("table_count") or 0)
    image_count = int(docx_structure.get("image_count") or 0)

    if paragraph_count < 10:
        failures.append(f"Word 段落数量异常偏少：{paragraph_count} < 10")
    if source_heading_count > 0 and heading_count == 0:
        failures.append("Word 中没有可识别标题样式，标题结构可能未正确写入。")
    elif source_heading_count > 0 and heading_count < max(1, source_heading_count // 2):
        warnings.append(f"Word 标题数量明显少于 Markdown 标题：{heading_count} < {source_heading_count}")

    if figure_count > 0 and image_count == 0:
        failures.append("figure_index.json 有图片计划，但 Word 中没有图片。")
    elif image_count < figure_count:
        warnings.append(f"Word 图片数量少于 figure_index.json：{image_count} < {figure_count}")

    if table_count > 0 and docx_table_count == 0:
        failures.append("table_index.json 有表格计划，但 Word 中没有表格。")
    elif docx_table_count < table_count:
        warnings.append(f"Word 表格数量少于 table_index.json：{docx_table_count} < {table_count}")

    return failures, warnings


def strict_visual_qa_failures(
    docx_structure: dict[str, Any],
    source_heading_count: int,
    source_content_chars: int,
    figure_count: int,
    table_count: int,
) -> tuple[list[str], list[str]]:
    failures, warnings = visual_qa_failures(docx_structure, source_heading_count, figure_count, table_count)
    if not docx_structure.get("exists") or docx_structure.get("error"):
        return failures, warnings

    heading_count = int(docx_structure.get("heading_count") or 0)
    docx_table_count = int(docx_structure.get("table_count") or 0)
    empty_table_count = int(docx_structure.get("empty_table_count") or 0)
    image_rel_count = int(docx_structure.get("image_count") or 0)
    inline_shape_count = int(docx_structure.get("inline_shape_count") or 0)
    docx_text_chars = int(docx_structure.get("nonspace_text_chars") or 0)

    if docx_text_chars < 200:
        failures.append(f"DOCX text payload is too small: {docx_text_chars} < 200")
    elif source_content_chars >= 1000 and docx_text_chars < max(200, source_content_chars // 3):
        failures.append(f"DOCX text payload is much smaller than source markdown: {docx_text_chars} < {source_content_chars // 3}")
    elif source_content_chars >= 1000 and docx_text_chars < source_content_chars // 2:
        warnings.append(f"DOCX text payload is smaller than source markdown: {docx_text_chars} < {source_content_chars // 2}")

    if source_heading_count > 0 and heading_count == 0:
        failures.append("No Word heading styles were detected while markdown headings exist")

    if figure_count > 0 and inline_shape_count == 0:
        failures.append(f"figure_index has figures but DOCX has no inline images: expected {figure_count}")
    elif inline_shape_count < figure_count:
        warnings.append(f"DOCX inline image count is lower than figure_index: {inline_shape_count} < {figure_count}")
    if image_rel_count > inline_shape_count:
        warnings.append(f"DOCX has image relationships not used as inline shapes: rels={image_rel_count}, inline={inline_shape_count}")

    if table_count > 0 and docx_table_count == 0:
        failures.append(f"table_index has tables but DOCX has no tables: expected {table_count}")
    if empty_table_count > 0:
        failures.append(f"DOCX contains empty tables: {empty_table_count}")

    return failures, warnings


def find_libreoffice() -> Path | None:
    candidates: list[Path] = []
    configured = str(__import__("os").environ.get("LIBREOFFICE_PATH") or "").strip()
    if configured:
        candidates.append(Path(configured))
    for executable in ("soffice", "libreoffice"):
        located = shutil.which(executable)
        if located:
            candidates.append(Path(located))
    candidates.extend(
        [
            Path(r"C:\Program Files\LibreOffice\program\soffice.exe"),
            Path(r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"),
            Path("/usr/bin/libreoffice"),
            Path("/usr/bin/soffice"),
            Path("/Applications/LibreOffice.app/Contents/MacOS/soffice"),
        ]
    )
    for candidate in candidates:
        if candidate.exists() and candidate.is_file():
            return candidate.resolve()
    return None


def render_docx_qa(path: Path, mode: str, source_content_chars: int) -> tuple[dict[str, Any], list[str], list[str]]:
    report: dict[str, Any] = {
        "mode": mode,
        "status": "SKIPPED" if mode == "skip" else "PENDING",
        "libreoffice": "",
        "pdf": "",
        "page_count": 0,
        "extracted_text_chars": 0,
        "returncode": None,
        "stdout": "",
    }
    failures: list[str] = []
    warnings: list[str] = []
    if mode == "skip":
        return report, failures, warnings
    if not path.exists():
        message = f"无法执行渲染 QA，Word 文件不存在：{rel(path)}"
        if mode == "required":
            failures.append(message)
            report["status"] = "FAIL"
        else:
            warnings.append(message)
            report["status"] = "UNAVAILABLE"
        return report, failures, warnings

    soffice = find_libreoffice()
    if soffice is None:
        message = "未找到 LibreOffice/soffice，无法执行 DOCX -> PDF 渲染检查。"
        if mode == "required":
            failures.append(message)
            report["status"] = "FAIL"
        else:
            warnings.append(message)
            report["status"] = "UNAVAILABLE"
        return report, failures, warnings

    report["libreoffice"] = str(soffice)
    RENDER_DIR.mkdir(parents=True, exist_ok=True)
    pdf_path = RENDER_DIR / f"{path.stem}.pdf"
    if pdf_path.exists():
        pdf_path.unlink()

    with tempfile.TemporaryDirectory(prefix="lo-profile-") as profile_dir:
        profile_uri = Path(profile_dir).resolve().as_uri()
        command = [
            str(soffice),
            f"-env:UserInstallation={profile_uri}",
            "--headless",
            "--nologo",
            "--nodefault",
            "--nofirststartwizard",
            "--convert-to",
            "pdf",
            "--outdir",
            str(RENDER_DIR.resolve()),
            str(path.resolve()),
        ]
        try:
            result = subprocess.run(
                command,
                cwd=str(BASE_DIR),
                text=True,
                encoding="utf-8",
                errors="replace",
                stdout=subprocess.PIPE,
                stderr=subprocess.STDOUT,
                timeout=120,
                check=False,
            )
            report["returncode"] = result.returncode
            report["stdout"] = result.stdout[-2000:]
        except Exception as exc:
            failures.append(f"LibreOffice 渲染命令异常：{type(exc).__name__}: {exc}")
            report["status"] = "FAIL"
            return report, failures, warnings

    if report["returncode"] != 0 or not pdf_path.exists() or pdf_path.stat().st_size <= 0:
        failures.append(
            f"LibreOffice 未生成有效 PDF：returncode={report['returncode']}，output={rel(pdf_path)}"
        )
        report["status"] = "FAIL"
        return report, failures, warnings

    report["pdf"] = rel(pdf_path)
    report["pdf_sha256"] = sha256_file(pdf_path)
    report["pdf_bytes"] = pdf_path.stat().st_size
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(pdf_path))
        page_count = len(reader.pages)
        extracted = "\n".join((page.extract_text() or "") for page in reader.pages)
        extracted_chars = len(re.sub(r"\s+", "", extracted))
        report["page_count"] = page_count
        report["extracted_text_chars"] = extracted_chars
        report["text_preview"] = extracted[:240]
    except Exception as exc:
        failures.append(f"渲染 PDF 无法读取：{type(exc).__name__}: {exc}")
        report["status"] = "FAIL"
        return report, failures, warnings

    if report["page_count"] < 1:
        failures.append("渲染 PDF 页数为 0。")
    minimum_rendered_text = 20 if source_content_chars < 500 else min(200, max(50, source_content_chars // 50))
    if report["extracted_text_chars"] < minimum_rendered_text:
        failures.append(
            f"渲染 PDF 可提取文本异常偏少：{report['extracted_text_chars']} < {minimum_rendered_text}"
        )
    report["status"] = "PASS" if not failures else "FAIL"
    return report, failures, warnings


def evaluate(render_mode: str = "auto") -> dict[str, Any]:
    source = source_path()
    outline = load_json(OUTLINE_FILE)
    figure_index = load_json(FIGURE_INDEX_FILE)
    table_index = load_json(TABLE_INDEX_FILE)
    failures: list[str] = []
    warnings: list[str] = []

    if not source.exists():
        failures.append(f"缺少正式论文源文件：{rel(SOURCE_FILE)}")
        text = ""
    else:
        text = source.read_text(encoding="utf-8")

    failures.extend(control_character_failures(text))
    failures.extend(internal_language_failures(text))
    duplicate_failures, duplicate_matches = duplicate_paragraph_failures(text)
    failures.extend(duplicate_failures)
    citation_qa, citation_failures, citation_warnings = citation_quality(text)
    failures.extend(citation_failures)
    warnings.extend(citation_warnings)

    counts = char_count(text)
    target_words = outline.get("target_words", {}) if isinstance(outline, dict) else {}
    min_words = int(target_words.get("min", 10000) or 10000)
    max_words = int(target_words.get("max", 22000) or 22000)
    if counts["content"] < min_words:
        failures.append(f"正文有效字数不足：{counts['content']} < {min_words}")
    if counts["content"] > max_words:
        warnings.append(f"正文有效字数超过建议上限：{counts['content']} > {max_words}")

    missing_sections = [label for label in REQUIRED_SECTIONS if not has_required_section(text, label)]
    for label in missing_sections:
        failures.append(f"缺少正式论文结构：{label}")

    if not re.search(r"(^|\n)\s*#*\s*5\.1\s+", text):
        failures.append("缺少 5.1 问题一模型章节。")
    if not re.search(r"(^|\n)\s*#*\s*5\.1\.1\s+", text):
        failures.append("缺少 5.1.1 三级标题。")
    if not re.search(r"(^|\n)\s*#*\s*5\.1\.2\s+", text):
        failures.append("缺少 5.1.2 三级标题。")

    question_reports = []
    for index, qid in enumerate(qids_from_outline(outline), start=1):
        q_failures: list[str] = []
        section = f"5.{index}"
        section_pattern = rf"(^|\n)\s*#*\s*{re.escape(section)}\s+"
        if not re.search(section_pattern, text):
            q_failures.append(f"缺少 {section} 对应 {qid} 的模型章节")
        for suffix, title in (
            ("1", "建模思路"),
            ("2", "变量定义与公式推导"),
            ("3", "求解算法"),
            ("4", "结果分析"),
            ("5", "模型检验或灵敏度分析"),
        ):
            if not re.search(rf"(^|\n)\s*#*\s*{re.escape(section + '.' + suffix)}\s+.*{title}", text):
                q_failures.append(f"缺少 {section}.{suffix} {title}")
        if not re.search(rf"{qid}|问题[一二三四五六七八九十{index}]", text):
            q_failures.append(f"正文未明确回扣 {qid}")
        if not re.search(rf"{section}[\s\S]{{0,5000}}Step\s*1", text, flags=re.IGNORECASE):
            q_failures.append(f"{section} 缺少 Step 1/Step 2 形式的算法步骤")
        question_reports.append({"question_id": qid, "status": "FAIL" if q_failures else "PASS", "failures": q_failures})
        failures.extend(q_failures)

    figures = index_items(figure_index, "figures", "figure_id")
    tables = index_items(table_index, "tables", "table_id")
    paper_tables = [item for item in tables if item.get("include_in_paper") is not False]
    missing_figures = [item.get("figure_id") for item in figures if not referenced(text, item, "figure_id")]
    missing_tables = [item.get("table_id") for item in tables if not referenced(text, item, "table_id")]
    for figure_id in missing_figures:
        failures.append(f"figure_index.json 中的图片未在正文引用：{figure_id}")
    for table_id in missing_tables:
        failures.append(f"table_index.json 中的表格未在正文引用：{table_id}")
    if len(figures) < 5:
        warnings.append(f"图数量少于展示样例建议值：{len(figures)} < 5")
    if len(tables) < 5:
        warnings.append(f"表数量少于展示样例建议值：{len(tables)} < 5")

    for placeholder in PLACEHOLDERS:
        if placeholder in text:
            failures.append(f"存在占位符或待补文本：{placeholder}")
    template_placeholders = re.findall(r"\{\{\s*[^{}\n]{1,80}\s*\}\}", text)
    for placeholder in template_placeholders[:12]:
        failures.append(f"存在模板占位符：{placeholder}")

    docx_structure = check_docx_structure(DOCX_FILE)
    if not docx_structure.get("exists"):
        failures.append(f"缺少正式 Word 文件：{rel(DOCX_FILE)}")
    elif docx_structure.get("error"):
        failures.append(f"Word 文件无法读取：{docx_structure['error']}")

    source_heading_count = markdown_heading_count(text)
    visual_failures, visual_warnings = strict_visual_qa_failures(
        docx_structure,
        source_heading_count,
        counts["content"],
        len(figures),
        len(paper_tables),
    )
    failures.extend(visual_failures)
    warnings.extend(visual_warnings)

    formula_tokens = source_formula_tokens(text)
    source_formula_count = len(formula_tokens)
    source_display_formula_count = sum(1 for item in formula_tokens if item.display)
    native_omml_count = int(docx_structure.get("native_omml_count") or 0)
    formula_failures: list[str] = []
    formula_warnings: list[str] = []
    if source_formula_count > native_omml_count:
        formula_failures.append(
            f"Word 原生公式数量不足：{native_omml_count} < {source_formula_count}。禁止用普通文本冒充公式。"
        )
    elif native_omml_count > source_formula_count:
        formula_warnings.append(
            f"Word 原生公式数量多于当前 Markdown 公式：{native_omml_count} > {source_formula_count}，请确认 Word 是否过期。"
        )
    failures.extend(formula_failures)
    warnings.extend(formula_warnings)

    render_qa, render_failures, render_warnings = render_docx_qa(DOCX_FILE, render_mode, counts["content"])
    failures.extend(render_failures)
    warnings.extend(render_warnings)

    input_paths = [
        source,
        DOCX_FILE,
        OUTLINE_FILE,
        FIGURE_INDEX_FILE,
        TABLE_INDEX_FILE,
        EVIDENCE_GATE_REPORT,
    ]
    input_hashes = {
        rel(path): sha256_file(path)
        for path in input_paths
        if path.exists() and path.is_file()
    }

    return {
        "schema_version": "1.0",
        "generated_by": "paper-formal-writer/scripts/check_paper_format.py",
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "status": "PASS" if not failures else "FAIL",
        "source": rel(source),
        "docx": rel(DOCX_FILE),
        "counts": counts,
        "target_words": {"min": min_words, "max": max_words},
        "question_reports": question_reports,
        "figure_count": len(figures),
        "table_count": len(tables),
        "paper_table_count": len(paper_tables),
        "missing_figures": missing_figures,
        "missing_tables": missing_tables,
        "source_heading_count": source_heading_count,
        "input_hashes": input_hashes,
        "citation_qa": citation_qa,
        "duplicate_paragraphs": duplicate_matches,
        "formula_qa": {
            "status": "PASS" if not formula_failures else "FAIL",
            "source_formula_count": source_formula_count,
            "source_display_formula_count": source_display_formula_count,
            "native_omml_count": native_omml_count,
            "native_display_omml_count": int(docx_structure.get("display_omml_count") or 0),
            "failures": formula_failures,
            "warnings": formula_warnings,
        },
        "docx_structure": docx_structure,
        "visual_qa": {
            "status": "PASS" if not visual_failures else "FAIL",
            "failures": visual_failures,
            "warnings": visual_warnings,
        },
        "render_qa": render_qa,
        "failures": failures,
        "warnings": warnings,
    }


def write_reports(report: dict[str, Any]) -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    REPORT_JSON.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    lines = [
        "# Formal Paper Format Check Report",
        "",
        f"- Status: `{report['status']}`",
        f"- Generated at: `{report['generated_at']}`",
        f"- Source: `{report['source']}`",
        f"- DOCX: `{report['docx']}`",
        f"- Effective chars: `{report['counts']['content']}`",
        f"- CJK chars: `{report['counts']['cjk']}`",
        f"- Figures in index: `{report['figure_count']}`",
        f"- Tables in index: `{report['table_count']}`",
        f"- Tables marked for paper: `{report['paper_table_count']}`",
        f"- Source formulas: `{report['formula_qa']['source_formula_count']}`",
        f"- Native Word equations: `{report['formula_qa']['native_omml_count']}`",
        f"- Body citations: `{report['citation_qa']['body_citation_count']}`",
        f"- Render QA: `{report['render_qa']['status']}`",
        "",
    ]
    if report["failures"]:
        lines.append("## Failures")
        lines.extend(f"- {item}" for item in report["failures"])
        lines.append("")
    if report["warnings"]:
        lines.append("## Warnings")
        lines.extend(f"- {item}" for item in report["warnings"])
        lines.append("")
    lines.append("## Question Coverage")
    for item in report["question_reports"]:
        lines.append(f"- {item['question_id']}: `{item['status']}`")
        for failure in item["failures"]:
            lines.append(f"  - {failure}")
    lines.append("")
    lines.append("## DOCX Structure")
    for key, value in report["docx_structure"].items():
        lines.append(f"- {key}: `{value}`")
    lines.append("")
    lines.append("## Visual QA")
    lines.append(f"- status: `{report['visual_qa']['status']}`")
    lines.append(f"- source_heading_count: `{report['source_heading_count']}`")
    for failure in report["visual_qa"]["failures"]:
        lines.append(f"- failure: {failure}")
    for warning in report["visual_qa"]["warnings"]:
        lines.append(f"- warning: {warning}")
    lines.append("")
    lines.append("## Formula QA")
    for key, value in report["formula_qa"].items():
        lines.append(f"- {key}: `{value}`")
    lines.append("")
    lines.append("## Citation QA")
    for key, value in report["citation_qa"].items():
        lines.append(f"- {key}: `{value}`")
    lines.append("")
    lines.append("## Render QA")
    for key, value in report["render_qa"].items():
        lines.append(f"- {key}: `{value}`")
    REPORT_MD.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")


def main() -> int:
    configure_utf8_stdio()
    parser = argparse.ArgumentParser(description="Check formal paper structure, native equations and rendered output.")
    parser.add_argument(
        "--render",
        choices=("auto", "required", "skip"),
        default="auto",
        help="auto renders when LibreOffice is available; required fails without a valid rendered PDF; skip disables render QA.",
    )
    args = parser.parse_args()
    report = evaluate(render_mode=args.render)
    write_reports(report)
    print(f"正式论文格式检查报告：{rel(REPORT_MD)}")
    if report["status"] == "PASS":
        print("✅ 正式论文格式门禁通过。")
        return 0
    print("⚠️ 正式论文格式门禁未通过。")
    for failure in report["failures"][:12]:
        print(f" - {failure}")
    if len(report["failures"]) > 12:
        print(f" - 其余 {len(report['failures']) - 12} 项见报告。")
    return 1


if __name__ == "__main__":
    raise SystemExit(main())
