from __future__ import annotations

import hashlib
import json
import math
import re
import sys
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt, Inches
from lite_common import safe_path, paper_checks


ROOT = Path.cwd()
OUTPUT_DIR = ROOT / "paper_output_lite"
MANIFEST_FILE = OUTPUT_DIR / "input_manifest.json"
PLAN_FILE = OUTPUT_DIR / "plan.json"
MODEL_FILE = OUTPUT_DIR / "code" / "model.py"
RUN_FILE = OUTPUT_DIR / "run_manifest.json"
RESULTS_FILE = OUTPUT_DIR / "results.json"
PAPER_FILE = OUTPUT_DIR / "paper.md"
DOCX_FILE = OUTPUT_DIR / "paper.docx"
REPORT_FILE = OUTPUT_DIR / "lite_report.json"

PLACEHOLDERS = (
    "待补",
    "待填写",
    "示例结果",
    "假设已经运行",
    "to be filled",
    "placeholder",
    "todo",
)
REQUIRED_SECTIONS = ("摘要", "问题重述", "模型", "结果", "结论")


def load_json(path: Path) -> Any:
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return None


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def resolve(path_text: object) -> Path:
    return safe_path(ROOT, path_text)


def rel(path: Path) -> str:
    try:
        return path.resolve().relative_to(ROOT.resolve()).as_posix()
    except ValueError:
        return path.resolve().as_posix()


def validate_records(records: object, label: str) -> list[str]:
    failures: list[str] = []
    if not isinstance(records, list):
        return [f"{label} 缺少文件记录。"]
    for record in records:
        if not isinstance(record, dict):
            failures.append(f"{label} 包含无效记录。")
            continue
        path = resolve(record.get("path"))
        if not path.exists():
            failures.append(f"{label}文件缺失：{rel(path)}")
            continue
        if path.stat().st_size != record.get("bytes") or sha256_file(path) != record.get("sha256"):
            failures.append(f"{label}文件已变化：{rel(path)}")
    return failures


def validate_metrics(value: object, qid: str) -> list[str]:
    failures: list[str] = []
    if value is None:
        return failures
    if not isinstance(value, dict):
        return [f"{qid} metrics 必须是对象。"]
    for name, metric in value.items():
        if not isinstance(metric, (int, float)) or isinstance(metric, bool) or not math.isfinite(float(metric)):
            failures.append(f"{qid} 指标 {name} 不是有限数值。")
    return failures


def clean_markdown(text: str) -> str:
    text = re.sub(r"!\[([^]]*)\]\([^)]+\)", r"[图：\1]", text)
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]*)`", r"\1", text)
    return text.strip()


def build_docx(text: str) -> None:
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"].font.size = Pt(11)
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            continue
        image = re.fullmatch(r"!\[([^]]*)\]\(([^)]+)\)", line)
        if image:
            document.add_picture(str(resolve(image.group(2))), width=Inches(6))
            if image.group(1):
                document.add_paragraph(image.group(1))
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            document.add_heading(clean_markdown(heading.group(2)), level=min(3, len(heading.group(1))))
        else:
            document.add_paragraph(clean_markdown(line))
    document.save(DOCX_FILE)
    reopened = Document(DOCX_FILE)
    if [p.text for p in reopened.paragraphs] != [p.text for p in document.paragraphs] or len(reopened.inline_shapes) != len(document.inline_shapes):
        raise ValueError("DOCX failed the reopen/content check")


def evaluate() -> dict[str, Any]:
    failures: list[str] = []
    warnings: list[str] = [
        "Lite 不检查原生 Word 公式、正文文献引用或 LibreOffice 渲染；需要这些能力时改用 Standard。"
    ]
    manifest = load_json(MANIFEST_FILE)
    plan = load_json(PLAN_FILE)
    run = load_json(RUN_FILE)
    results = load_json(RESULTS_FILE)

    if not isinstance(manifest, dict) or manifest.get("status") != "PASS":
        failures.append("缺少通过的 input_manifest.json。")
    else:
        failures.extend(validate_records(manifest.get("files"), "输入"))
        from lite_run import input_failures
        failures.extend(input_failures(manifest))

    questions = plan.get("questions") if isinstance(plan, dict) else None
    if not isinstance(questions, list) or not questions:
        failures.append("plan.json 缺少非空 questions 列表。")
        questions = []
    plan_ids = [str(item.get("id") or "").strip() for item in questions if isinstance(item, dict)]
    if len(plan_ids) != len(questions) or any(not qid for qid in plan_ids) or len(set(plan_ids)) != len(plan_ids):
        failures.append("plan.json 的问题 ID 为空或重复。")

    if not isinstance(run, dict) or run.get("status") != "PASS":
        failures.append("缺少通过的 run_manifest.json；请运行 lite_run.py。")
    else:
        if not MODEL_FILE.exists() or sha256_file(MODEL_FILE) != run.get("script_sha256"):
            failures.append("model.py 在运行后已变化；请重新运行 lite_run.py。")
        failures.extend(validate_records(run.get("inputs"), "运行输入"))
        failures.extend(validate_records(run.get("outputs"), "运行输出"))
        expected = {rel(p): sha256_file(p) for p in (MODEL_FILE, PLAN_FILE, MANIFEST_FILE) if p.is_file()}
        if run.get("initial_hashes") != expected or len(expected) != 3:
            failures.append("plan, script or input manifest changed since the run; rerun lite_run.py")
        if not any(r.get("path") == rel(RESULTS_FILE) for r in run.get("outputs", [])):
            failures.append("run manifest does not bind results.json")

    result_questions = results.get("questions") if isinstance(results, dict) else None
    if not isinstance(results, dict) or results.get("status") != "computed":
        failures.append("results.json status 必须为 computed。")
    if not isinstance(result_questions, list) or not result_questions:
        failures.append("results.json 缺少非空 questions 列表。")
        result_questions = []
    result_map = {
        str(item.get("id") or "").strip(): item
        for item in result_questions
        if isinstance(item, dict) and str(item.get("id") or "").strip()
    }
    if len(result_map) != len(result_questions) or set(result_map) != set(plan_ids):
        failures.append("results must cover exactly the planned question IDs without duplicates")
    for qid in plan_ids:
        item = result_map.get(qid)
        if not item:
            failures.append(f"results.json 缺少 {qid}。")
            continue
        if not str(item.get("answer") or "").strip():
            failures.append(f"{qid} 缺少非空 answer。")
        failures.extend(validate_metrics(item.get("metrics"), qid))
        evidence = item.get("evidence", [])
        if evidence is not None and not isinstance(evidence, list):
            failures.append(f"{qid} evidence 必须是列表。")
        for path_text in evidence if isinstance(evidence, list) else []:
            path = resolve(path_text)
            if not path.exists() or not path.is_file() or path.stat().st_size == 0:
                failures.append(f"{qid} 证据文件缺失或为空：{rel(path)}")

    paper_text = PAPER_FILE.read_text(encoding="utf-8") if PAPER_FILE.exists() else ""
    if not paper_text.strip():
        failures.append("缺少非空 paper_output_lite/paper.md。")
    lowered = paper_text.lower()
    for placeholder in PLACEHOLDERS:
        if placeholder.lower() in lowered:
            failures.append(f"paper.md 包含占位文本：{placeholder}")
    for section in REQUIRED_SECTIONS:
        if not any(section in h for h in re.findall(r"(?m)^\s*#{1,6}\s+[^\n]+", paper_text)):
            failures.append(f"paper.md 缺少必要章节关键词：{section}")
    for qid in plan_ids:
        if qid not in paper_text:
            failures.append(f"paper.md 未明确覆盖 {qid}。")
    if isinstance(plan, dict):
        failures.extend(paper_checks(paper_text, plan, result_map))
    recorded_outputs = {item.get("path") for item in run.get("outputs", [])} if isinstance(run, dict) else set()
    for _, path_text in re.findall(r"!\[([^]]*)\]\(([^)]+)\)", paper_text):
        path = resolve(path_text)
        if rel(path) not in recorded_outputs or not path.is_file():
            failures.append(f"paper image is not a recorded model output: {path_text}")

    status = "PASS" if not failures else "FAIL"
    return {
        "schema_version": "1.0",
        "generated_by": "mathmodel-lite/scripts/lite_finalize.py",
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "status": status,
        "failures": failures,
        "warnings": warnings,
        "question_ids": plan_ids,
        "acceptance_scope": "NOT_ACCEPTED" if failures else ("LITE_BASIC_REPORT_ONLY" if ((plan or {}).get("delivery") or {}).get("mode", "basic-report") == "basic-report" else "LITE_SHORT_OR_TEST_ONLY"),
        "input_hashes": {
            rel(path): sha256_file(path)
            for path in (MANIFEST_FILE, PLAN_FILE, MODEL_FILE, RUN_FILE, RESULTS_FILE, PAPER_FILE)
            if path.exists() and path.is_file()
        },
    }


def main() -> int:
    from lite_common import configure_stdio
    configure_stdio()
    try:
        safe_path(ROOT, "paper_output_lite")
        for path in (MANIFEST_FILE, PLAN_FILE, MODEL_FILE, RUN_FILE, RESULTS_FILE, PAPER_FILE, DOCX_FILE, REPORT_FILE):
            safe_path(ROOT, path.relative_to(ROOT).as_posix())
    except ValueError as exc:
        print(f"[FAIL] {exc}")
        return 1
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    try:
        report = evaluate()
        if report["status"] == "PASS":
            build_docx(PAPER_FILE.read_text(encoding="utf-8"))
            report["docx"] = {"path": rel(DOCX_FILE), "bytes": DOCX_FILE.stat().st_size, "sha256": sha256_file(DOCX_FILE)}
    except Exception as exc:
        report = {"status": "FAIL", "acceptance_scope": "NOT_ACCEPTED", "failures": [f"invalid artifact or Word generation failed: {exc}"]}
    REPORT_FILE.write_text(json.dumps(report, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(f"Lite report: {rel(REPORT_FILE)}")
    if report["status"] == "PASS":
        print(f"[PASS] {report['acceptance_scope']}: basic Word report, not competition-format acceptance.")
        return 0
    for failure in report["failures"]:
        print(f"[FAIL] {failure}")
    return 1


if __name__ == "__main__":
    sys.exit(main())

